'''
if you need more information please you can contact me imadmaalouf02@gmail.com or 0642781650 
just send me a message in email or whatsapp

'''

'''
please read the readme file for more information about how to use the tool and the features of the tool
and olso read what i wirte here in the code as a comment this is important information that will help you to use the tool and understand it better
'''

'''
You can use the lighning ia to run all the content in the notebook you can create a simple account and the free version is enough ghtning.ai
and you can run the code using the lighning ia just  dowload the code and upload it to the lighning ia and run it there
you don't need to install anything in your local machine and olso in the lighning ia you don't need to install streamlit 
it is already installed defaulltly in the lighning framework you need just to open tge studio and go to the + in the right 
of the window and click on the + and then click on the streamlit  and  it's done
but if you use your local machine you need to install the requirements in the requirements.txt file
and olso you need to install streamlit in your local machine
you can install it using the command pip install streamlit
and then you can run the code using the command streamlit run main_app_LAST_V.py    

'''



import streamlit as st
import pandas as pd
import re
import io
from pathlib import Path
from typing import List, Dict, Optional, Tuple, Set
import zipfile
from xml.etree import ElementTree as ET
from dataclasses import dataclass
import tempfile
import base64
import plotly.express as px
import plotly.graph_objects as go
from collections import Counter, defaultdict
import docx
from docx.document import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

@dataclass
class Requirement:
    """Enhanced data class for storing requirement information"""
    id: str
    description: str
    page: int
    section: str
    table_index: int = 0
    row_index: int = 0
    confidence: float = 1.0

class EnhancedPageTracker:
    """Advanced page tracking with multiple strategies"""
    
    def __init__(self):
        self.estimated_chars_per_page = 2500
        self.estimated_lines_per_page = 45
        self.page_break_markers = [
            'page break', 'new page', 'next page',
            'début de page', 'nouvelle page'
        ]
    
    def track_pages_comprehensive(self, document: Document) -> Dict[int, int]:
        """Map element indices to page numbers using multiple strategies"""
        element_to_page = {}
        current_page = 1
        cumulative_chars = 0
        element_index = 0
        
        for element in document.element.body:
            # Strategy 1: Look for explicit page breaks
            if self._has_page_break(element):
                current_page += 1
                cumulative_chars = 0
            
            # Strategy 2: Estimate based on content length
            element_text = self._get_element_text(element)
            element_chars = len(element_text)
            
            if cumulative_chars + element_chars > self.estimated_chars_per_page:
                current_page += 1
                cumulative_chars = element_chars
            else:
                cumulative_chars += element_chars
            
            element_to_page[element_index] = current_page
            element_index += 1
        
        return element_to_page
    
    def _has_page_break(self, element) -> bool:
        """Check if element contains page break indicators"""
        try:
            # Check for w:br with type="page"
            for br in element.iter():
                if br.tag.endswith('}br'):
                    type_attr = br.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type')
                    if type_attr == 'page':
                        return True
            
            # Check for section breaks
            for sect in element.iter():
                if sect.tag.endswith('}sectPr'):
                    return True
                    
        except Exception:
            pass
        return False
    
    def _get_element_text(self, element) -> str:
        """Extract text content from XML element"""
        try:
            text_parts = []
            for text_elem in element.iter():
                if text_elem.tag.endswith('}t'):
                    if text_elem.text:
                        text_parts.append(text_elem.text)
            return ' '.join(text_parts)
        except Exception:
            return ""

class SectionAwareExtractor:
    """Enhanced extractor with section awareness and accurate page mapping"""
    
    def __init__(self):
        self.page_tracker = EnhancedPageTracker()
        self.section_patterns = [
            r'(\d+\.)*\d+\.?\d*\s*[Ee]xigences?\s+[Ff]onctionnelles?',
            r'(\d+\.)*\d+\.?\d*\s*[Ff]unctional\s+[Rr]equirements?',
            r'(\d+\.)*\d+\.?\d*\s*[Rr]equirements?\s+[Ff]onctionnels?',
            r'(\d+\.)*\d+\.?\d*\s*[Nn]on[-\s]*[Ff]unctional\s+[Rr]equirements?',
            r'(\d+\.)*\d+\.?\d*\s*[Rr]equirements?\s+[Tt]echniques?',
            r'(\d+\.)*\d+\.?\d*\s*[Ss]pécifications?\s+[Tt]echniques?'
        ]
        
        self.requirement_id_patterns = {
            'WAVE5': r'WAVE5[-_][A-Z]+[-_][A-Z]+[-_][A-Z]+[-_]\d+(?:$ \d+ $ )?',
            'Generic Automotive': r'[A-Z]{2,8}[-_]\d{3,6}',
            'Hierarchical': r'[A-Z]+[-_]\d+[-_]\d+[-_]\d+',
            'Standard Format': r'[A-Z]+[-_][A-Z]+[-_]\d+',
            'Extended WAVE': r'WAVE\d+[-_][A-Z]{3,8}[-_][A-Z]{2,8}[-_][A-Z]{3,8}[-_]\d+',
            'Custom Pattern': r'[A-Z0-9]+[-_][A-Z0-9]+[-_][A-Z0-9]+[-_]\d+'
        }
    
    def extract_requirements_enhanced(
        self, 
        document: Document, 
        selected_patterns: List[str],
        start_page: int = 1,
        end_page: int = 999,
        section_filter: Optional[str] = None
    ) -> List[Requirement]:
        """Enhanced extraction with section awareness and page filtering"""
        
        requirements = []
        element_to_page = self.page_tracker.track_pages_comprehensive(document)
        current_section = "Unknown Section"
        element_index = 0
        
        # Process document elements in order
        for element in document.element.body:
            current_page = element_to_page.get(element_index, 1)
            
            # Skip if outside page range
            if current_page < start_page or current_page > end_page:
                element_index += 1
                continue
            
            # Check for section headers
            section_title = self._extract_section_title(element)
            if section_title:
                if not section_filter or self._section_matches_filter(section_title, section_filter):
                    current_section = section_title
            
            # Process tables for requirements
            if element.tag.endswith('}tbl'):
                table_requirements = self._extract_from_xml_table(
                    element, selected_patterns, current_page, current_section, element_index
                )
                requirements.extend(table_requirements)
            
            element_index += 1
        
        # Also process docx document object for additional context
        for i, paragraph in enumerate(document.paragraphs):
            estimated_page = min(i // 20 + 1, end_page)  # Rough estimate
            if start_page <= estimated_page <= end_page:
                section_title = self._extract_section_from_paragraph(paragraph)
                if section_title:
                    if not section_filter or self._section_matches_filter(section_title, section_filter):
                        current_section = section_title
        
        for table_idx, table in enumerate(document.tables):
            estimated_page = min(table_idx * 2 + 1, end_page)  # Rough estimate
            if start_page <= estimated_page <= end_page:
                if not section_filter or self._section_matches_filter(current_section, section_filter):
                    table_reqs = self._extract_from_docx_table(
                        table, selected_patterns, estimated_page, current_section, table_idx
                    )
                    # Avoid duplicates by checking if requirement already exists
                    for req in table_reqs:
                        if not any(r.id == req.id and r.description[:50] == req.description[:50] for r in requirements):
                            requirements.append(req)
        
        return self._deduplicate_requirements(requirements)
    
    def _extract_element_text(self, element) -> str:
        """Extract text content from XML element - FIXED METHOD"""
        try:
            text_parts = []
            for text_elem in element.iter():
                if text_elem.tag.endswith('}t'):
                    if text_elem.text:
                        text_parts.append(text_elem.text)
            return ' '.join(text_parts).strip()
        except Exception:
            return ""
    
    def _extract_section_title(self, element) -> Optional[str]:
        """Extract section title from XML element"""
        text = self._extract_element_text(element)
        
        if not text or len(text) > 200:  # Section titles are usually short
            return None
        
        for pattern in self.section_patterns:
            if re.search(pattern, text, re.IGNORECASE):
                return text.strip()
        
        # Additional heuristics for section detection
        if (re.match(r'^\d+\.', text) and 
            any(keyword in text.lower() for keyword in ['requirement', 'exigence', 'functional', 'technique'])):
            return text.strip()
        
        return None
    
    def _extract_section_from_paragraph(self, paragraph: Paragraph) -> Optional[str]:
        """Extract section title from docx paragraph"""
        text = paragraph.text.strip()
        
        if not text or len(text) > 200:
            return None
        
        # Check if paragraph is a heading style
        if paragraph.style.name.startswith('Heading'):
            return text
        
        # Check for section patterns
        for pattern in self.section_patterns:
            if re.search(pattern, text, re.IGNORECASE):
                return text
        
        return None
    
    def _section_matches_filter(self, section: str, section_filter: str) -> bool:
        """Check if section matches filter criteria"""
        if not section_filter or section_filter.lower() == "all":
            return True
        
        filter_lower = section_filter.lower()
        section_lower = section.lower()
        
        # Direct match
        if filter_lower in section_lower:
            return True
        
        # Keyword-based matching
        filter_keywords = {
            'functional': ['functional', 'fonctionnel'],
            'non-functional': ['non-functional', 'non functional', 'non-fonctionnel'],
            'technical': ['technical', 'technique'],
            'performance': ['performance', 'performant'],
            'safety': ['safety', 'sécurité', 'securite']
        }
        
        if filter_lower in filter_keywords:
            return any(keyword in section_lower for keyword in filter_keywords[filter_lower])
        
        return False
    
    def _extract_from_xml_table(
        self, 
        table_element, 
        selected_patterns: List[str],
        page: int,
        section: str,
        table_index: int
    ) -> List[Requirement]:
        """Extract requirements from XML table element"""
        requirements = []
        
        try:
            rows = table_element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tr')
            
            for row_idx, row in enumerate(rows):
                cells = row.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tc')
                
                if len(cells) >= 2:
                    # Extract text from first two cells
                    cell1_text = self._extract_cell_text_from_xml(cells[0])
                    cell2_text = self._extract_cell_text_from_xml(cells[1])
                    
                    # Try to find requirement ID in either cell
                    req_id = self._extract_requirement_id(cell1_text, selected_patterns)
                    description = cell2_text
                    
                    if not req_id:
                        req_id = self._extract_requirement_id(cell2_text, selected_patterns)
                        description = cell1_text
                    
                    if req_id and description and len(description.strip()) > 10:
                        confidence = self._calculate_confidence(req_id, description)
                        
                        requirements.append(Requirement(
                            id=req_id,
                            description=self._clean_description(description),
                            page=page,
                            section=section,
                            table_index=table_index,
                            row_index=row_idx,
                            confidence=confidence
                        ))
        
        except Exception as e:
            st.warning(f"Error processing XML table: {e}")
        
        return requirements
    
    def _extract_cell_text_from_xml(self, cell_element) -> str:
        """Extract text from XML table cell"""
        try:
            text_parts = []
            for text_elem in cell_element.iter():
                if text_elem.tag.endswith('}t'):
                    if text_elem.text:
                        text_parts.append(text_elem.text)
            return ' '.join(text_parts).strip()
        except Exception:
            return ""
    
    def _extract_from_docx_table(
        self,
        table: Table,
        selected_patterns: List[str],
        page: int,
        section: str,
        table_index: int
    ) -> List[Requirement]:
        """Extract requirements from docx table object"""
        requirements = []
        
        try:
            for row_idx, row in enumerate(table.rows):
                if len(row.cells) >= 2:
                    cell1_text = row.cells[0].text.strip()
                    cell2_text = row.cells[1].text.strip()
                    
                    # Try to find requirement ID
                    req_id = self._extract_requirement_id(cell1_text, selected_patterns)
                    description = cell2_text
                    
                    if not req_id and len(row.cells) > 2:
                        req_id = self._extract_requirement_id(cell2_text, selected_patterns)
                        description = row.cells[2].text.strip() if len(row.cells) > 2 else cell1_text
                    
                    if req_id and description and len(description.strip()) > 10:
                        confidence = self._calculate_confidence(req_id, description)
                        
                        requirements.append(Requirement(
                            id=req_id,
                            description=self._clean_description(description),
                            page=page,
                            section=section,
                            table_index=table_index,
                            row_index=row_idx,
                            confidence=confidence
                        ))
        
        except Exception as e:
            st.warning(f"Error processing docx table: {e}")
        
        return requirements
    
    def _extract_requirement_id(self, text: str, selected_patterns: List[str]) -> Optional[str]:
        """Extract requirement ID using selected patterns"""
        if not text:
            return None
        
        for pattern_name in selected_patterns:
            if pattern_name in self.requirement_id_patterns:
                pattern = self.requirement_id_patterns[pattern_name]
                match = re.search(pattern, text, re.IGNORECASE)
                if match:
                    return match.group(0)
        
        return None
    
    def _clean_description(self, description: str) -> str:
        """Clean and normalize description text"""
        # Remove excessive whitespace
        description = re.sub(r'\s+', ' ', description)
        
        # Remove common table artifacts
        description = re.sub(r'^(Description|Libellé|Requirement|Exigence)[:\s]*', '', description, flags=re.IGNORECASE)
        
        # Remove trailing artifacts
        description = re.sub(r'\s*(Input requirement|Exigence amont).*$', '', description, flags=re.IGNORECASE)
        
        return description.strip()
    
    def _calculate_confidence(self, req_id: str, description: str) -> float:
        """Calculate confidence score for requirement extraction"""
        confidence = 1.0
        
        # Reduce confidence for very short descriptions
        if len(description) < 20:
            confidence *= 0.7
        
        # Reduce confidence if description looks like a header
        if any(word in description.lower() for word in ['description', 'requirement', 'exigence', 'libellé']):
            confidence *= 0.8
        
        # Increase confidence for well-structured IDs
        if re.match(r'^[A-Z]+[-_][A-Z]+[-_][A-Z]+[-_]\d+', req_id):
            confidence *= 1.1
        
        return min(confidence, 1.0)
    
    def _deduplicate_requirements(self, requirements: List[Requirement]) -> List[Requirement]:
        """Remove duplicate requirements based on ID and description similarity"""
        seen_ids = set()
        unique_requirements = []
        
        for req in requirements:
            # Create a key for deduplication
            key = f"{req.id}_{req.description[:100]}"
            
            if key not in seen_ids:
                seen_ids.add(key)
                unique_requirements.append(req)
        
        return unique_requirements

def main():
    st.set_page_config(
        page_title="🚗 Enhanced Automotive Requirements Extractor",
        page_icon="🚗",
        layout="wide"
    )
    
    st.title("🚗 Enhanced Automotive Requirements Extractor")
    st.markdown("*Extract functional requirements from Word documents with precise page mapping and section awareness*")
    
    # Initialize extractor
    extractor = SectionAwareExtractor()
    
    # Sidebar controls
    with st.sidebar:
        st.header("📋 Extraction Settings")
        
        # File upload
        uploaded_file = st.file_uploader(
            "Upload Word Document",
            type=['docx'],
            help="Upload .docx files containing automotive requirements"
        )
        
        # Page range settings
        st.subheader("📄 Page Range")
        col1, col2 = st.columns(2)
        with col1:
            start_page = st.number_input("Start Page", min_value=1, value=1)
        with col2:
            end_page = st.number_input("End Page", min_value=1, value=999)
        
        # Pattern selection
        st.subheader("🔍 ID Patterns")
        available_patterns = list(extractor.requirement_id_patterns.keys())
        selected_patterns = st.multiselect(
            "Select Requirement ID Patterns",
            available_patterns,
            default=['WAVE5', 'Generic Automotive'],
            help="Choose patterns that match your document's requirement ID format"
        )
        
        # Section filtering
        st.subheader("🏷️ Section Filter")
        section_filter = st.selectbox(
            "Filter by Section",
            ["All", "Functional", "Non-Functional", "Technical", "Performance", "Safety"],
            help="Extract requirements only from specific sections"
        )
        
        # Advanced settings
        with st.expander("⚙️ Advanced Settings"):
            min_confidence = st.slider(
                "Minimum Confidence",
                min_value=0.0,
                max_value=1.0,
                value=0.5,
                step=0.1,
                help="Filter requirements by extraction confidence"
            )
            
            show_duplicates = st.checkbox(
                "Show Potential Duplicates",
                value=False,
                help="Include requirements that might be duplicates"
            )
    
    # Main processing area
    if uploaded_file and selected_patterns:
        try:
            # Load document
            document = docx.Document(uploaded_file)
            
            # Show document info
            st.info(f"📄 Processing document: **{uploaded_file.name}** | Pages: {start_page}-{end_page} | Patterns: {', '.join(selected_patterns)}")
            
            # Extract requirements
            with st.spinner("🔄 Extracting requirements with enhanced section awareness..."):
                requirements = extractor.extract_requirements_enhanced(
                    document=document,
                    selected_patterns=selected_patterns,
                    start_page=start_page,
                    end_page=end_page,
                    section_filter=section_filter if section_filter != "All" else None
                )
            
            # Filter by confidence
            filtered_requirements = [r for r in requirements if r.confidence >= min_confidence]
            
            if filtered_requirements:
                st.success(f"✅ Extracted {len(filtered_requirements)} requirements from {len(set(r.section for r in filtered_requirements))} sections")
                
                # Convert to DataFrame
                df_data = []
                for req in filtered_requirements:
                    df_data.append({
                        'ID': req.id,
                        'Description': req.description,
                        'Page': req.page,
                        'Section': req.section,
                        'Table Index': req.table_index,
                        'Row Index': req.row_index,
                        'Confidence': round(req.confidence, 2)
                    })
                
                df = pd.DataFrame(df_data)
                
                # Create tabs for different views
                tab1, tab2, tab3, tab4 = st.tabs(["📊 Results", "📈 Analytics", "🔍 Details", "💾 Export"])
                
                with tab1:
                    st.subheader("📋 Extracted Requirements")
                    
                    # Section filter for display
                    available_sections = ['All'] + list(df['Section'].unique())
                    display_section = st.selectbox(
                        "Display Section:",
                        available_sections,
                        key="display_section"
                    )
                    
                    display_df = df if display_section == 'All' else df[df['Section'] == display_section]
                    
                    # Display with formatting
                    st.dataframe(
                        display_df,
                        use_container_width=True,
                        column_config={
                            'Description': st.column_config.TextColumn(width="large"),
                            'Confidence': st.column_config.ProgressColumn(min_value=0, max_value=1),
                            'Page': st.column_config.NumberColumn(format="%d")
                        }
                    )
                
                with tab2:
                    st.subheader("📈 Extraction Analytics")
                    
                    col1, col2 = st.columns(2)
                    
                    with col1:
                        # Section distribution
                        section_counts = df['Section'].value_counts()
                        fig_sections = px.pie(
                            values=section_counts.values,
                            names=section_counts.index,
                            title="Requirements by Section"
                        )
                        st.plotly_chart(fig_sections, use_container_width=True)
                    
                    with col2:
                        # Page distribution
                        page_counts = df['Page'].value_counts().sort_index()
                        fig_pages = px.bar(
                            x=page_counts.index,
                            y=page_counts.values,
                            title="Requirements by Page",
                            labels={'x': 'Page', 'y': 'Count'}
                        )
                        st.plotly_chart(fig_pages, use_container_width=True)
                    
                    # Summary statistics
                    col3, col4, col5, col6 = st.columns(4)
                    with col3:
                        st.metric("Total Requirements", len(df))
                    with col4:
                        st.metric("Unique Sections", df['Section'].nunique())
                    with col5:
                        st.metric("Page Range", f"{df['Page'].min()}-{df['Page'].max()}")
                    with col6:
                        st.metric("Avg Confidence", f"{df['Confidence'].mean():.2f}")
                
                with tab3:
                    st.subheader("🔍 Detailed Information")
                    
                    # Pattern analysis
                    st.write("**Pattern Distribution:**")
                    pattern_stats = {}
                    for pattern_name in selected_patterns:
                        pattern = extractor.requirement_id_patterns[pattern_name]
                        count = sum(1 for req in filtered_requirements if re.match(pattern, req.id, re.IGNORECASE))
                        pattern_stats[pattern_name] = count
                    
                    pattern_df = pd.DataFrame(list(pattern_stats.items()), columns=['Pattern', 'Count'])
                    st.dataframe(pattern_df, use_container_width=True)
                    
                    # Confidence distribution
                    st.write("**Confidence Distribution:**")
                    fig_confidence = px.histogram(
                        df,
                        x='Confidence',
                        nbins=20,
                        title="Confidence Score Distribution"
                    )
                    st.plotly_chart(fig_confidence, use_container_width=True)
                
                with tab4:
                    st.subheader("💾 Export Options")
                    
                    col1, col2 = st.columns(2)
                    
                    with col1:
                        # CSV export
                        csv = df.to_csv(index=False)
                        st.download_button(
                            label="📥 Download as CSV",
                            data=csv,
                            file_name=f"requirements_{uploaded_file.name.split('.')[0]}_{start_page}_{end_page}.csv",
                            mime="text/csv",
                            use_container_width=True
                        )
                    
                    with col2:
                        # Excel export with multiple sheets
                        output = io.BytesIO()
                        with pd.ExcelWriter(output, engine='openpyxl') as writer:
                            # Main sheet
                            df.to_excel(writer, sheet_name='All Requirements', index=False)
                            
                            # Section-based sheets
                            for section in df['Section'].unique():
                                section_df = df[df['Section'] == section]
                                safe_name = re.sub(r'[^\w\s-]', '', section)[:31]
                                section_df.to_excel(writer, sheet_name=safe_name, index=False)
                        
                        st.download_button(
                            label="📊 Download as Excel",
                            data=output.getvalue(),
                            file_name=f"requirements_{uploaded_file.name.split('.')[0]}_{start_page}_{end_page}.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                            use_container_width=True
                        )
                    
                    # Export summary
                    st.info(f"""
                    **Export Summary:**
                    - Total Requirements: {len(df)}
                    - Sections: {df['Section'].nunique()}
                    - Page Range: {df['Page'].min()}-{df['Page'].max()}
                    - Excel includes separate sheets per section
                    """)
            
            else:
                st.warning("⚠️ No requirements found matching the specified criteria. Try adjusting the patterns or confidence threshold.")
        
        except Exception as e:
            st.error(f"❌ Error processing document: {e}")
            with st.expander("🐛 Debug Information"):
                import traceback
                st.code(traceback.format_exc())
    
    elif uploaded_file and not selected_patterns:
        st.warning("⚠️ Please select at least one requirement ID pattern.")
    
    else:
        # Welcome screen
        st.markdown("""
        ## 🎯 Key Features
        
        - **📄 Accurate Page Mapping**: Precise page number tracking using advanced heuristics
        - **📋 Section-Aware Extraction**: Identify and group requirements by document sections
        - **🔍 Pattern-Based Filtering**: Multiple automotive requirement ID formats
        - **📊 Table Structure Detection**: Extract from structured requirement tables
        - **📈 Advanced Analytics**: Visualize distribution across pages and sections
        
        ## 📖 How to Use
        
        1. **Upload** your automotive Word document (.docx)
        2. **Set page range** to focus extraction on specific pages
        3. **Choose ID patterns** that match your document format
        4. **Filter by section** to extract only relevant requirements
        5. **Review** structured results with page numbers and sections
        6. **Export** to Excel with section-based worksheets
        
        ## 📋 Example Input Structure
        
        ```
        6.1.1.5 Functional Requirements / Exigences fonctionnelles
        
        | Requirement ID         | Description              | Status |
        |------------------------|--------------------------|---------|
        | WAVE5-AVAS-ST-FUNC-001 | IF sound is required... | Active |
        ```
        
        ## 📊 Enhanced Output
        
        | ID | Description | Page | Section | Source |
        |----|-------------|------|---------|---------|
        | WAVE5-AVAS-ST-FUNC-001 | IF sound is required AND speed ≤ max THEN emit sound | 15 | Functional Requirements | Table 1, Row 2 |
        """)

if __name__ == "__main__":
    main()
